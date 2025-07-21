# file_generator.py

from docx import Document
from fpdf import FPDF
import io

def create_docx(report_text):
    """Creates a DOCX file in memory from the report text."""
    document = Document()
    document.add_heading('Cybersecurity Incident Report', 0)
    for paragraph in report_text.split('\n'):
        # Add paragraphs, attempting to handle simple markdown like headings
        if paragraph.startswith('## '):
            document.add_heading(paragraph.lstrip('## '), level=2)
        elif paragraph.startswith('# '):
            document.add_heading(paragraph.lstrip('# '), level=1)
        else:
            document.add_paragraph(paragraph)
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_pdf(report_text):
    """Creates a PDF file in memory from the report text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    cleaned_text = report_text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, txt=cleaned_text)

    return bytes(pdf.output())